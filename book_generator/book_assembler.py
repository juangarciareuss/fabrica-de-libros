# book_generator/book_assembler.py

from docx import Document
import pypandoc
import logging
import os
import re

## --- FUNCIÓN PRINCIPAL (ORQUESTADOR) ---
def create_final_document(book_title, book_content, curated_sources, output_folder, template_path, include_internal_citations=False):
    """
    Crea una versión específica del documento de Word (limpia o con citas internas).
    Esta es la función central que orquesta la creación del documento.
    """
    # Intenta usar una plantilla si se proporciona, si no, crea un documento en blanco.
    try:
        doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
        if template_path and os.path.exists(template_path):
            logging.info(f"Usando plantilla: '{template_path}'")
    except Exception as e:
        logging.warning(f"No se pudo usar la plantilla de Word. Se creará un documento en blanco. Error: {e}")
        doc = Document()

    doc.add_heading(book_title, level=0)
    
    # Este set registrará los IDs numéricos de las fuentes que se citen en el texto.
    cited_sources_map = set()

    # --- LÓGICA DE PROCESAMIENTO DE CAPÍTULOS ---
    introduction = next((chap for chap in book_content if chap.get('type') == 'introduction'), None)
    conclusion = next((chap for chap in book_content if chap.get('type') == 'conclusion'), None)
    main_chapters = [chap for chap in book_content if chap.get('type') not in ['introduction', 'conclusion']]

    # Añadir capítulos al documento
    if introduction:
        _add_chapter_to_doc(doc, introduction['title'], introduction['content'], cited_sources_map, include_internal_citations)

    for chapter in main_chapters:
        doc.add_page_break()
        _add_chapter_to_doc(doc, chapter['title'], chapter['content'], cited_sources_map, include_internal_citations)

    if conclusion:
        doc.add_page_break()
        _add_chapter_to_doc(doc, conclusion['title'], conclusion['content'], cited_sources_map, include_internal_citations)

    # --- SECCIÓN DE BIBLIOGRAFÍA ---
    _add_bibliography_to_doc(doc, curated_sources, cited_sources_map)
    
    # --- GUARDADO DEL ARCHIVO ---
    suffix = "_CON_REFERENCIAS" if include_internal_citations else "_FINAL_LIMPIO"
    doc_path = os.path.join(output_folder, f"{book_title}{suffix}.docx")
    doc.save(doc_path)
    logging.info(f"Documento '{os.path.basename(doc_path)}' guardado.")
    return doc_path


## --- FUNCIONES DE AYUDA (ESPECIALISTAS) ---

def _add_chapter_to_doc(document, title, content, cited_sources_map, include_internal_citations):
    """Añade un único capítulo formateado al documento."""
    document.add_heading(title, level=1)
    
    CITE_PATTERN = r'\[CITA:\s*([\d,\s]+)\]'
    
    if not isinstance(content, str):
        content = str(content)

    # Limpieza de marcadores de 'Fragmento' que la IA pudiera dejar
    content = re.sub(r'\[`?Fragmento\s*#\d+`?\]|\(`?Fragmento\s*#\d+`?\)', '', content)

    paragraphs = content.split('\n')
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Buscamos las citas para añadirlas al mapa de la bibliografía
        matches = re.findall(CITE_PATTERN, para)
        if matches:
            for id_group in matches:
                ids = [int(s.strip()) for s in id_group.split(',') if s.strip().isdigit()]
                cited_sources_map.update(ids)

        # Si es la versión final para el lector, eliminamos los marcadores [CITA: X]
        final_para = para
        if not include_internal_citations:
            final_para = re.sub(CITE_PATTERN, '', para).strip()
        
        if final_para:
            document.add_paragraph(final_para)

def _add_bibliography_to_doc(document, curated_sources, cited_sources_map):
    """Añade la sección de bibliografía formateada al final del documento."""
    document.add_page_break()
    document.add_heading("Bibliografía y Fuentes", level=1)
    
    if not cited_sources_map or not curated_sources:
        document.add_paragraph("No se citaron fuentes externas en este documento.")
        return

    # Creamos un diccionario para buscar fácilmente los detalles de las fuentes por su ID
    sources_dict = {src.get('id'): src for src in curated_sources if 'id' in src}
    
    # Ordenamos los IDs citados numéricamente para que la bibliografía esté ordenada
    sorted_cited_ids = sorted(list(cited_sources_map))
    
    for source_id in sorted_cited_ids:
        source_details = sources_dict.get(source_id)
        p = document.add_paragraph()
        p.add_run(f"[{source_id}] ").bold = True
        if source_details:
            source_url = source_details.get('url', source_details.get('source', 'URL no disponible'))
            snippet_text = source_details.get('snippet', 'Contenido no disponible')
            
            p.add_run(f"Fuente: {source_url}\n")
            p.add_run(f"Fragmento: ").italic = True
            p.add_run(f"\"{snippet_text}\"").italic = True
        else:
            p.add_run(f"Detalles para la fuente con ID {source_id} no encontrados.")

def convert_to_epub(docx_path, book_title, output_folder="output"):
    """Convierte el documento de Word (versión limpia) a formato EPUB."""
    if not docx_path or not os.path.exists(docx_path): 
        logging.error(f"No se encontró el archivo DOCX '{docx_path}' para la conversión a EPUB.")
        return None
    
    epub_path = os.path.join(output_folder, f"{book_title}.epub")
    try:
        logging.info("Iniciando conversión a EPUB...")
        pypandoc.convert_file(docx_path, 'epub', outputfile=epub_path, extra_args=['--toc'])
        logging.info(f"Archivo EPUB guardado en: {epub_path}")
        return epub_path
    except Exception as e:
        logging.error(f"Error durante la conversión a EPUB: {e}")
        return None