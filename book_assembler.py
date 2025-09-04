# book_generator/book_assembler.py

from docx import Document
import pypandoc
import logging
import os
import re

def create_final_document(book_title, book_content, curated_sources, output_folder, template_path, include_internal_citations=False):
    """
    Crea una versión específica del documento de Word (limpia o con citas internas).
    """
    try:
        doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
    except Exception as e:
        logging.warning(f"No se pudo usar la plantilla de Word. Se creará un documento en blanco. Error: {e}")
        doc = Document()

    doc.add_heading(book_title, level=0)
    cited_sources_map = set()

    # --- Lógica de Procesamiento de Capítulos ---
    introduction = next((chap for chap in book_content if chap.get('type') == 'introduction'), None)
    conclusion = next((chap for chap in book_content if chap.get('type') == 'conclusion'), None)
    main_chapters = [chap for chap in book_content if chap.get('type') not in ['introduction', 'conclusion']]

    if introduction:
        _add_chapter_to_doc(doc, introduction['title'], introduction['content'], curated_sources, cited_sources_map, include_internal_citations)

    for chapter in main_chapters:
        doc.add_page_break()
        _add_chapter_to_doc(doc, chapter['title'], chapter['content'], curated_sources, cited_sources_map, include_internal_citations)

    if conclusion:
        doc.add_page_break()
        _add_chapter_to_doc(doc, conclusion['title'], conclusion['content'], curated_sources, cited_sources_map, include_internal_citations)

    _add_bibliography_to_doc(doc, curated_sources, cited_sources_map)
    
    suffix = "_CON_REFERENCIAS" if include_internal_citations else "_FINAL_LIMPIO"
    doc_path = os.path.join(output_folder, f"{book_title}{suffix}.docx")
    
    try:
        doc.save(doc_path)
        logging.info(f"Documento '{os.path.basename(doc_path)}' guardado.")
        return doc_path
    except PermissionError:
        logging.critical(f"¡ERROR DE PERMISOS AL GUARDAR EL ARCHIVO '{doc_path}'!")
        logging.critical("CAUSA MÁS PROBABLE: El archivo está abierto en otro programa (ej. Microsoft Word).")
        return None

def _add_chapter_to_doc(document, title, content, cited_sources_map, include_internal_citations):
    """Añade un único capítulo formateado al documento, manejando citas de forma robusta."""
    document.add_heading(title, level=1)
    
    # Expresión regular mejorada para capturar CUALQUIER contenido dentro de [CITA: ...]
    CITE_PATTERN = r'\[CITA:\s*([^\]]+)\]'
    
    if not isinstance(content, str):
        content = str(content)

    content = re.sub(r'\[`?Fragmento\s*#\d+`?\]|\(`?Fragmento\s*#\d+`?\)', '', content)
    paragraphs = content.split('\n')

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Registrar todas las citas encontradas
        matches = re.findall(CITE_PATTERN, para)
        for match in matches:
            # Limpiamos y dividimos por si hay múltiples citas en un solo bloque
            source_keys = [s.strip() for s in match.split(',')]
            for key in source_keys:
                cited_sources_map.add(key) # Añadimos la clave (sea ID o URL) al set

        final_para = para
        if not include_internal_citations:
            final_para = re.sub(CITE_PATTERN, '', para).strip()
        
        if final_para:
            document.add_paragraph(final_para)

def _add_bibliography_to_doc(document, curated_sources, cited_sources_map):
    """Añade una bibliografía completa y formateada al final del documento."""
    document.add_page_break()
    document.add_heading("Bibliografía y Fuentes", level=1)
    
    if not cited_sources_map or not curated_sources:
        document.add_paragraph("No se citaron fuentes externas en este documento.")
        return

    # Crear diccionarios para búsqueda rápida por ID y por URL
    sources_by_id = {str(src.get('id')): src for src in curated_sources if 'id' in src}
    sources_by_url = {src.get('url'): src for src in curated_sources if 'url' in src}
    
    # Usamos una lista para poder ordenarla y mostrarla
    processed_sources = []
    
    for key in sorted(list(cited_sources_map)):
        source_details = sources_by_id.get(key) or sources_by_url.get(key)
        if source_details and source_details not in processed_sources:
            processed_sources.append(source_details)

    for i, source in enumerate(processed_sources, 1):
        p = document.add_paragraph()
        p.add_run(f"[{i}] ").bold = True
        source_url = source.get('url', 'URL no disponible')
        snippet_text = source.get('snippet', 'Contenido no disponible')
        
        p.add_run(f"Fuente: {source_url}\n")
        p.add_run(f"Fragmento: ").italic = True
        p.add_run(f"\"{snippet_text}\"").italic = True


def convert_to_epub(docx_path, book_title, output_folder="output"):
    # (Esta función no cambia)
    if not docx_path: return None
    epub_path = os.path.join(output_folder, f"{book_title}.epub")
    try:
        logging.info("Iniciando conversión a EPUB...")
        pypandoc.convert_file(docx_path, 'epub', outputfile=epub_path, extra_args=['--toc'])
        logging.info(f"Archivo EPUB guardado en: {epub_path}")
        return epub_path
    except Exception as e:
        logging.error(f"Error durante la conversión a EPUB: {e}")
        return None

