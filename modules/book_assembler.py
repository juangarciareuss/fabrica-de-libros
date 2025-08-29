from docx import Document
import pypandoc
import logging
import os

def create_word_document(book_title, book_content, template_path="MiPlantilla.docx", output_folder="output"):
    """Crea un documento de Word usando una plantilla de estilos predefinida."""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    doc_path = os.path.join(output_folder, f"{book_title}.docx")

    # --- CAMBIO CLAVE: Cargar la plantilla ---
    try:
        document = Document(template_path)
        logging.info(f"Plantilla '{template_path}' cargada exitosamente.")
    except Exception as e:
        logging.warning(f"No se pudo encontrar la plantilla '{template_path}'. Se usará un documento por defecto. Error: {e}")
        document = Document()

    # Limpiar el contenido preexistente de la plantilla (si lo hubiera)
    for para in document.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # Añadir contenido nuevo usando los estilos de la plantilla
    document.add_heading(book_title, level=0) # Usa el estilo 'Title' de tu plantilla

    for chapter in book_content:
        logging.info(f"Añadiendo capítulo '{chapter['title']}' al documento.")
        # Aquí aplicamos los estilos que DEBEN existir en tu plantilla
        try:
            document.add_heading(chapter['title'], level=1) # Usa el estilo 'Heading 1'
            
            content = chapter.get('content', '')
            if content:
                paragraphs = content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        # Usa el estilo por defecto para el cuerpo de texto ('Normal' o 'Body Text')
                        document.add_paragraph(para)
            
            document.add_page_break()
        except Exception as e:
            logging.error(f"Error al aplicar estilo. Asegúrate de que los estilos 'Heading 1', 'Title', etc., existen en tu plantilla. Error: {e}")

    try:
        document.save(doc_path)
        logging.info(f"Documento de Word guardado exitosamente en: {doc_path}")
        return doc_path
    except Exception as e:
        logging.error(f"Error al guardar el documento de Word: {e}")
        return None

def convert_to_epub(docx_path, book_title, output_folder="output"):
    if not docx_path:
        logging.error("No se proporcionó ruta de DOCX para la conversión.")
        return None
    epub_path = os.path.join(output_folder, f"{book_title}.epub")
    try:
        logging.info("Iniciando conversión a EPUB con Pandoc...")
        pypandoc.convert_file(docx_path, 'epub', outputfile=epub_path, extra_args=['--toc'])
        logging.info(f"Archivo EPUB guardado exitosamente en: {epub_path}")
        return epub_path
    except Exception as e:
        logging.error(f"Error durante la conversión a EPUB. Asegúrate de que Pandoc esté instalado. Error: {e}")
        return None